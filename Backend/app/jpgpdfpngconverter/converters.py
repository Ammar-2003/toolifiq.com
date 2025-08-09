import os
import io
from django.conf import settings
from django.core.exceptions import ValidationError
from PIL import Image
import fitz  # PyMuPDF
from .models import FileConversion
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pdf2docx import Converter
import logging
from threading import Lock
from concurrent.futures import ThreadPoolExecutor
import shutil
import subprocess
from pathlib import Path
import tempfile
from pdfminer.high_level import extract_pages
from pdfminer.layout import LAParams
from pdfminer.high_level import extract_text_to_fp
from pdfminer.converter import HTMLConverter
from bs4 import BeautifulSoup
from html import escape as HTML  # If you're using HTML escaping
import pdfplumber
import pandas as pd
from django.core.files.storage import default_storage


logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)



class FileConverter:
    def __init__(self, file_obj, conversion_type):
        self.file_obj = file_obj
        self.conversion_type = conversion_type
        self.conversion = None
    
    def validate_file(self):
        """Validate file based on conversion type with strict checks"""
        if self.conversion_type == 'jpg2pdf':
            if not self.file_obj.name.lower().endswith(('.jpg', '.jpeg')):
                raise ValidationError('Only JPG files are allowed for JPG to PDF conversion')
            
            try:
                with Image.open(self.file_obj) as img:
                    img.verify()
                self.file_obj.seek(0)
            except Exception:
                raise ValidationError('Invalid JPG image file')
                
        elif self.conversion_type == 'pdf2jpg':
            if not self.file_obj.name.lower().endswith('.pdf'):
                raise ValidationError('Only PDF files are allowed for PDF to JPG conversion')
            
            pdf_header = self.file_obj.read(4)
            self.file_obj.seek(0)
            if pdf_header != b'%PDF':
                raise ValidationError('Invalid PDF file format')
        
        elif self.conversion_type == 'png2pdf':
            if not self.file_obj.name.lower().endswith('.png'):
                raise ValidationError('Only PNG files are allowed for PNG to PDF conversion')
            try:
                with Image.open(self.file_obj) as img:
                    if img.format != 'PNG':
                        raise ValidationError('Invalid PNG file')
                self.file_obj.seek(0)
            except Exception:
                raise ValidationError('Invalid PNG image file')
                
        elif self.conversion_type == 'pdf2png':
            if not self.file_obj.name.lower().endswith('.pdf'):
                raise ValidationError('Only PDF files are allowed for PDF to PNG conversion')
            pdf_header = self.file_obj.read(4)
            self.file_obj.seek(0)
            if pdf_header != b'%PDF':
                raise ValidationError('Invalid PDF file format')

    def get_output_filename(self, extension):
        """Generate output filename preserving the original name but changing extension"""
        original_name = os.path.splitext(self.file_obj.name)[0]
        return f"{original_name}.{extension}"
    
    def create_conversion_record(self):
        self.validate_file()
        self.conversion = FileConversion.objects.create(
            original_file=self.file_obj,
            conversion_type=self.conversion_type
        )
        return self.conversion
    
    def get_output_path(self, extension):
        """Get absolute filesystem path for output"""
        output_filename = self.get_output_filename(extension)
        return os.path.join(
            settings.MEDIA_ROOT,
            'converted_files',
            f"{self.conversion.id}_{output_filename}"
        )
    
    def save_conversion(self, output_path, file_extension):
        """Save conversion result to database"""
        output_filename = self.get_output_filename(file_extension)
        self.conversion.converted_file.name = f'converted_files/{self.conversion.id}_{output_filename}'
        self.conversion.save()
        return self.conversion

    def convert_pdf_to_jpg(self, pdf_path, output_path):
        """Convert PDF to JPG with optimized quality and basic compression

        Args:
            pdf_path: Path to input PDF file
            output_path: Path to save output JPG file

        Raises:
            ValueError: If conversion fails or produces invalid output
        """
        try:
            # Validate input file
            if not os.path.exists(pdf_path):
                raise ValueError("PDF file does not exist")

            # Open PDF document
            doc = fitz.open(pdf_path)

            if not doc.is_pdf:
                doc.close()
                raise ValueError("Input file is not a valid PDF")

            if len(doc) == 0:
                doc.close()
                raise ValueError("PDF document is empty")

            # Convert first page to image
            page = doc[0]
            zoom = 2
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            doc.close()

            # Convert to PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Save with optimized JPEG settings
            img.save(
                output_path,
                format='JPEG',
                quality=90,
                optimize=True,
                progressive=False,
                subsampling=0,
                dpi=(300, 300)
            )

            # Validate output
            if not os.path.exists(output_path):
                raise ValueError("Output file was not created")

            if os.path.getsize(output_path) == 0:
                os.remove(output_path)
                raise ValueError("Conversion produced empty file")

        except Exception as e:
            # Clean up if anything went wrong
            if 'doc' in locals() and doc:
                doc.close()
            if os.path.exists(output_path):
                try:
                    os.remove(output_path)
                except:
                    pass
            raise ValueError(f"PDF to JPG conversion failed: {str(e)}")

    def convert_jpg_to_pdf(self, jpg_paths, output_path, is_multiple=False):
        """Convert JPGs to PDF with enhanced quality and smart sizing:
        - Single file: Keep original dimensions with max quality
        - Multiple files: Standardize to consistent larger dimensions (20% larger than largest image)
        """
        try:
            QUALITY = 70        # Maximum quality
            MARGIN = 30         # 30pt margin (≈10.5mm)
            MIN_SIZE = 500      # Minimum dimension for small images
            
            if not is_multiple:
                with Image.open(jpg_paths) as img:
                    if img.mode != 'RGB':
                        img = img.convert('RGB')
                    img.save(
                        output_path,
                        format='PDF',
                        quality=QUALITY,
                        optimize=True,
                        dpi=(300, 300)
                    )
            else:
                pdf_pages = []
                max_width, max_height = 0, 0
                
                with Image.open(jpg_paths[0]) as first_img:
                    max_width, max_height = first_img.size
                
                for jpg_path in jpg_paths[1:]:
                    with Image.open(jpg_path) as img:
                        max_width = max(max_width, img.width)
                        max_height = max(max_height, img.height)
                
                # Calculate target dimensions (20% larger than largest image)
                target_width = min(int(max_width * 1.2), 1200)  # Cap at 1200px
                target_height = min(int(max_height * 1.2), 1600)  # Cap at 1600px
                
                for jpg_path in jpg_paths:
                    with Image.open(jpg_path) as img:
                        if img.mode != 'RGB':
                            img = img.convert('RGB')
                        
                        # Create canvas with target dimensions
                        new_img = Image.new('RGB', (target_width, target_height), (255, 255, 255))
                        
                        # Calculate available space with margins
                        paste_width = target_width - (2 * MARGIN)
                        paste_height = target_height - (2 * MARGIN)
                        
                        # Maintain aspect ratio while fitting within available space
                        img.thumbnail(
                            (paste_width, paste_height),
                            resample=Image.LANCZOS
                        )
                        
                        # Center the image
                        x = (target_width - img.width) // 2
                        y = (target_height - img.height) // 2
                        new_img.paste(img, (x, y))
                        
                        pdf_pages.append(new_img)
                
                if pdf_pages:
                    pdf_pages[0].save(
                        output_path,
                        format='PDF',
                        save_all=True,
                        append_images=pdf_pages[1:],
                        quality=QUALITY,
                        optimize=True,
                        dpi=(300, 300)
                    )
            
            if not os.path.exists(output_path):
                raise ValueError("Output file was not created")
            if os.path.getsize(output_path) == 0:
                raise ValueError("Conversion produced empty PDF file")
        
        except Exception as e:
            if os.path.exists(output_path):
                try:
                    os.remove(output_path)
                except OSError:
                    pass
            raise ValueError(f"JPG to PDF conversion failed: {str(e)}")


    def convert_png_to_pdf(self, png_paths, output_path, is_multiple=False):
        """Convert PNG(s) to PDF with professional handling:
        - Single file: Original dimensions with max quality
        - Multiple files: Standardized A4 pages with optimally sized images
        - Handles transparency by converting to white background
        - Maintains high quality with smart compression
        - Larger image size with reduced margins
        """
        try:
            # Page settings (A4 dimensions at 300dpi)
            A4_WIDTH = 2480
            A4_HEIGHT = 3508
            QUALITY = 95
            MIN_FILE_SIZE = 1024  # 1KB minimum
            
            if not is_multiple:
                # Single file conversion - preserve original size
                with Image.open(png_paths) as img:
                    # Handle transparency
                    if img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info):
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        background.paste(img, mask=img.split()[-1])
                        img = background
                    elif img.mode != 'RGB':
                        img = img.convert('RGB')
                    
                    img.save(
                        output_path,
                        format='PDF',
                        quality=QUALITY,
                        optimize=True,
                        dpi=(300, 300),
                        save_all=True,
                        compression='zip'
                    )
            else:
                # Multiple file conversion - standardized pages with larger images
                pdf_pages = []
                
                for png_path in png_paths:
                    with Image.open(png_path) as img:
                        # Handle transparency
                        if img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info):
                            background = Image.new('RGB', img.size, (255, 255, 255))
                            background.paste(img, mask=img.split()[-1])
                            img = background
                        elif img.mode != 'RGB':
                            img = img.convert('RGB')
                        
                        # Create new A4 page with white background
                        new_page = Image.new('RGB', (A4_WIDTH, A4_HEIGHT), (255, 255, 255))
                        
                        # Calculate original aspect ratio
                        original_ratio = img.width / img.height
                        page_ratio = A4_WIDTH / A4_HEIGHT
                        
                        # Determine optimal scaling based on orientation
                        if original_ratio > page_ratio:
                            # Landscape image - scale to full width
                            scale_factor = (A4_WIDTH * 0.9) / img.width  # 90% of page width
                        else:
                            # Portrait image - scale to full height
                            scale_factor = (A4_HEIGHT * 0.9) / img.height  # 90% of page height
                        
                        # Apply scaling
                        new_width = int(img.width * scale_factor)
                        new_height = int(img.height * scale_factor)
                        
                        # Resize with high-quality interpolation
                        img = img.resize((new_width, new_height), Image.LANCZOS)
                        
                        # Calculate centered position
                        x_offset = (A4_WIDTH - img.width) // 2
                        y_offset = (A4_HEIGHT - img.height) // 2
                        
                        # Paste image onto page
                        new_page.paste(img, (x_offset, y_offset))
                        pdf_pages.append(new_page)
                
                # Save all pages to PDF
                if pdf_pages:
                    pdf_pages[0].save(
                        output_path,
                        format='PDF',
                        save_all=True,
                        append_images=pdf_pages[1:],
                        quality=QUALITY,
                        optimize=True,
                        dpi=(300, 300),
                        compression='zip'
                    )
            
            # Validate output
            if not os.path.exists(output_path):
                raise ValueError("Output PDF was not created")
            if os.path.getsize(output_path) < MIN_FILE_SIZE:
                raise ValueError("Output PDF appears too small, conversion may have failed")
        
        except Exception as e:
            if os.path.exists(output_path):
                try:
                    os.remove(output_path)
                except OSError as cleanup_error:
                    raise ValueError(f"Conversion failed and cleanup also failed: {str(cleanup_error)}")
            raise ValueError(f"PNG to PDF conversion failed: {str(e)}")



    def convert_pdf_to_pngs(self, pdf_path, output_folder):
        """
        Convert PDF to multiple PNGs (one per page)
        Returns list of generated PNG file paths
        """
        png_files = []
        try:
            doc = fitz.open(pdf_path)
            if not doc.is_pdf:
                raise ValueError("Invalid PDF file")

            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Calculate DPI based on original PDF dimensions
                zoom = 96 / 72  # Standard 96 DPI
                mat = fitz.Matrix(zoom, zoom)

                pix = page.get_pixmap(
                    matrix=mat,
                    alpha=False,
                    colorspace="RGB",
                    dpi=96
                )

                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # Save each page as PNG
                png_path = os.path.join(output_folder, f"page_{page_num+1}.png")
                img.save(png_path, format='PNG', optimize=True, compress_level=6)
                png_files.append(png_path)

            doc.close()
            return png_files

        except Exception as e:
            # Cleanup any partial conversions
            for f in png_files:
                if os.path.exists(f):
                    os.remove(f)
            raise ValueError(f"PDF conversion failed: {str(e)}")
    
    def convert_pdf_to_webp(self, pdf_path, output_path):
        """Convert PDF to WebP with optimized quality and basic compression

        Args:
            pdf_path: Path to input PDF file
            output_path: Path to save output WebP file

        Raises:
            ValueError: If conversion fails or produces invalid output
        """
        try:
            # Validate input file
            if not os.path.exists(pdf_path):
                raise ValueError("PDF file does not exist")

            # Open PDF document
            doc = fitz.open(pdf_path)

            if not doc.is_pdf:
                doc.close()
                raise ValueError("Input file is not a valid PDF")

            if len(doc) == 0:
                doc.close()
                raise ValueError("PDF document is empty")

            # Convert first page to image
            page = doc[0]
            zoom = 2
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            doc.close()

            # Convert to PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Save with optimized WebP settings
            img.save(
                output_path,
                format='WEBP',
                quality=80,
                method=6
            )

            # Validate output
            if not os.path.exists(output_path):
                raise ValueError("Output file was not created")

            if os.path.getsize(output_path) == 0:
                os.remove(output_path)
                raise ValueError("Conversion produced empty file")

        except Exception as e:
            # Clean up if anything went wrong
            if 'doc' in locals() and doc:
                doc.close()
            if os.path.exists(output_path):
                try:
                    os.remove(output_path)
                except:
                    pass
            raise ValueError(f"PDF to WebP conversion failed: {str(e)}")
        
    

    def convert(self):
        """Main conversion method that routes to specific converters"""
        self.create_conversion_record()
        
        # Create temp input file path
        input_filename = os.path.join(
            settings.MEDIA_ROOT,
            'temp_uploads',
            f"{self.conversion.id}_{self.file_obj.name}"
        )
        
        # Ensure directory exists
        os.makedirs(os.path.dirname(input_filename), exist_ok=True)
        
        # Save the uploaded file temporarily
        with open(input_filename, 'wb+') as destination:
            for chunk in self.file_obj.chunks():
                destination.write(chunk)
        
        try:
            if self.conversion_type == 'jpg2pdf':
                output_path = self.get_output_path('pdf')
                self.convert_jpg_to_pdf(input_filename, output_path)
                return self.save_conversion(output_path, 'pdf')
            
            elif self.conversion_type == 'pdf2jpg':
                output_path = self.get_output_path('jpg')
                self.convert_pdf_to_jpg(input_filename, output_path)
                return self.save_conversion(output_path, 'jpg')
        
            elif self.conversion_type == 'pdf2webp':
                output_path = self.get_output_path('webp')
                self.convert_pdf_to_jpg(input_filename, output_path)
                return self.save_conversion(output_path, 'webp')
            
            elif self.conversion_type == 'png2pdf':
                output_path = self.get_output_path('pdf')
                self.convert_png_to_pdf(input_filename, output_path)
                return self.save_conversion(output_path, 'pdf')
            
            elif self.conversion_type == 'pdf2png':
                output_path = self.get_output_path('png')
                self.convert_pdf_to_png(input_filename, output_path)
                return self.save_conversion(output_path, 'png')
            
            else:
                raise ValueError(f"Unsupported conversion type: {self.conversion_type}")
                
        finally:
            # Clean up the temporary file
            if os.path.exists(input_filename):
                os.remove(input_filename)

class PdfToWordConverter:
    """Production-grade PDF to Word conversion service"""
    
    def __init__(self):
        self.executor = ThreadPoolExecutor(max_workers=4)
        self.lock = Lock()
    
    def _validate_pdf(self, pdf_path):
        """Validate PDF file integrity"""
        try:
            with fitz.open(pdf_path) as doc:
                if not doc.is_pdf:
                    raise ValueError("Invalid PDF file")
                if doc.needs_pass:
                    raise ValueError("Password-protected PDFs are not supported")
                return len(doc)
        except Exception as e:
            raise ValueError(f"PDF validation failed: {str(e)}")
    
    def _preserve_layout_conversion(self, pdf_path, output_path):
        """High-fidelity conversion preserving layout and graphics"""
        try:
            with self.lock:  # pdf2docx isn't fully thread-safe
                cv = Converter(pdf_path)
                cv.convert(output_path, start=0, end=None)
                cv.close()
            
            # Post-process to clean up common artifacts
            doc = Document(output_path)
            
            # Remove empty paragraphs
            for paragraph in list(doc.paragraphs):
                if not paragraph.text.strip():
                    p = paragraph._element
                    p.getparent().remove(p)
            
            # Ensure proper page breaks
            for i, section in enumerate(doc.sections):
                if i > 0:  # Skip first section
                    section.start_type
                    
            doc.save(output_path)
            return True
        except Exception as e:
            logger.error(f"Layout preservation failed: {str(e)}")
            if os.path.exists(output_path):
                os.remove(output_path)
            raise
    
    def _text_extraction_conversion(self, pdf_path, output_path):
        """Clean text extraction with basic formatting"""
        try:
            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            with fitz.open(pdf_path) as pdf_doc:
                for page in pdf_doc:
                    blocks = page.get_text("dict")["blocks"]
                    
                    # Add page break (except for first page)
                    if page.number > 0:
                        doc.add_page_break()
                    
                    for b in blocks:
                        if b["type"] == 0:  # Text block
                            paragraph = doc.add_paragraph()
                            
                            for l in b["lines"]:
                                for s in l["spans"]:
                                    run = paragraph.add_run(s["text"])
                                    
                                    # Preserve formatting
                                    if "bold" in s["font"].lower():
                                        run.bold = True
                                    if "italic" in s["font"].lower():
                                        run.italic = True
                                    
                                    # Preserve color if not black
                                    if s["color"] != 0:
                                        rgb = (
                                            s["color"] >> 16 & 0xff,
                                            s["color"] >> 8 & 0xff,
                                            s["color"] & 0xff
                                        )
                                        if rgb != (0, 0, 0):
                                            run.font.color.rgb = RGBColor(*rgb)
                                    
                                    # Preserve font size (clamped)
                                    run.font.size = Pt(max(8, min(36, s["size"] * 0.7)))
                            
                            # Preserve alignment
                            if "align" in b:
                                align_map = {
                                    0: WD_PARAGRAPH_ALIGNMENT.LEFT,
                                    1: WD_PARAGRAPH_ALIGNMENT.CENTER,
                                    2: WD_PARAGRAPH_ALIGNMENT.RIGHT,
                                    3: WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                }
                                paragraph.alignment = align_map.get(b["align"], WD_PARAGRAPH_ALIGNMENT.LEFT)
            
            doc.save(output_path)
            return True
        except Exception as e:
            logger.error(f"Text extraction failed: {str(e)}")
            if os.path.exists(output_path):
                os.remove(output_path)
            raise
    
    def convert_pdf_to_word(self, pdf_path, output_path, preserve_graphics=True):
        """Convert PDF to Word with production-grade quality
        
        Args:
            pdf_path: Path to input PDF
            output_path: Output DOCX path
            preserve_graphics: Whether to preserve layout (True) or extract text (False)
            
        Returns:
            dict: Conversion metadata including page count and file size
        """
        try:
            # Validate input
            page_count = self._validate_pdf(pdf_path)
            
            # Perform conversion
            if preserve_graphics:
                self._preserve_layout_conversion(pdf_path, output_path)
            else:
                self._text_extraction_conversion(pdf_path, output_path)
            
            # Validate output
            if not os.path.exists(output_path):
                raise ValueError("Conversion failed - no output file created")
            if os.path.getsize(output_path) == 0:
                os.remove(output_path)
                raise ValueError("Conversion produced empty file")
            
            return {
                "page_count": page_count,
                "file_size": os.path.getsize(output_path),
                "preserved_layout": preserve_graphics
            }
            
        except Exception as e:
            logger.error(f"PDF to Word conversion failed: {str(e)}")
            raise

class PdfConversionError(Exception):
    """Custom exception for PDF conversion failures"""
    pass


class PdfToHtmlConverter:
    """
    Professional PDF to HTML converter with multiple conversion strategies.
    Uses MinIO for temporary file storage in Docker environment.
    """
    
    def __init__(self, file_obj):
        self.file_obj = file_obj
        self.temp_dir = None
        self.conversion_methods = [
            self._convert_with_pymupdf,
            self._convert_with_pdfminer_enhanced,
            self._convert_with_pdfplumber
        ]
        
    def convert_to_formatted_html(self):
        """
        Convert PDF to HTML with best-effort layout preservation.
        
        Returns:
            Path: Path to the converted HTML file
            
        Raises:
            PdfConversionError: If conversion fails
        """
        try:
            self.temp_dir = Path(tempfile.mkdtemp())
            pdf_path = self._save_uploaded_file()
            output_path = self._get_output_path('html')
            
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Try conversion methods in order of quality
            for method in self.conversion_methods:
                try:
                    method(pdf_path, output_path)
                    if self._validate_output(output_path):
                        return output_path
                except Exception as e:
                    logger.warning(f"Method {method.__name__} failed: {str(e)}")
                    continue
            
            raise PdfConversionError("All conversion methods failed")
            
        except Exception as e:
            self.cleanup()
            raise PdfConversionError(f"Formatted conversion failed: {str(e)}")
    
    def convert_to_clean_text(self):
        """
        Convert PDF to clean, semantic HTML with minimal styling.
        
        Returns:
            Path: Path to the converted HTML file
            
        Raises:
            PdfConversionError: If conversion fails
        """
        try:
            self.temp_dir = Path(tempfile.mkdtemp())
            pdf_path = self._save_uploaded_file()
            output_path = self._get_output_path('html')
            
            # Extract structured text content using pdfplumber
            text_blocks = []
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    words = page.extract_words(
                        x_tolerance=2,
                        y_tolerance=2,
                        keep_blank_chars=False,
                        use_text_flow=True,
                        extra_attrs=["fontname", "size"]
                    )
                    
                    for word in words:
                        text_blocks.append({
                            'text': word['text'],
                            'x0': word['x0'],
                            'x1': word['x1'],
                            'y0': word['top'],
                            'y1': word['bottom'],
                            'page': page_num,
                            'font_size': word['size'],
                            'bold': 'bold' in word['fontname'].lower() if word['fontname'] else False,
                            'italic': 'italic' in word['fontname'].lower() if word['fontname'] else False
                        })
            
            # Generate clean HTML
            html_content = self._generate_clean_html(text_blocks)
            
            # Write output file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return output_path
            
        except Exception as e:
            self.cleanup()
            raise PdfConversionError(f"Clean text conversion failed: {str(e)}")
    
    def _convert_with_pdfminer_enhanced(self, pdf_path, output_path):
        """High-quality conversion using PDFMiner with enhanced layout analysis"""
        laparams = LAParams(
            line_overlap=0.5,
            line_margin=0.5,
            word_margin=0.1,
            boxes_flow=0.7,
            detect_vertical=True,
            all_texts=True,
            char_margin=2.0
        )
        
        with open(output_path, 'w', encoding='utf-8') as out_f:
            with open(pdf_path, 'rb') as pdf_f:
                converter = HTMLConverter(
                    out_f,
                    codec='utf-8',
                    laparams=laparams,
                    scale=1.0,
                    layoutmode='exact',
                    showpageno=False,
                    pagemargin=0,
                    imagewriter=None,
                    debug=False
                )
                
                for page in extract_pages(pdf_f):
                    converter.process_page(page)
                
                converter.close()
        
        self._post_process_html(output_path, preserve_layout=True)
    
    def _convert_with_pymupdf(self, pdf_path, output_path):
        """Conversion using PyMuPDF with precise layout preservation"""
        doc = fitz.open(pdf_path)
        
        html_content = []
        for page in doc:
            rect = page.rect
            width = rect.width
            height = rect.height
            
            blocks = page.get_text("dict")["blocks"]
            
            page_html = []
            for b in blocks:
                if "lines" in b:
                    block_html = []
                    for line in b["lines"]:
                        line_html = []
                        for span in line["spans"]:
                            style = ""
                            if "bold" in span["font"].lower():
                                style += "font-weight:bold;"
                            if "italic" in span["font"].lower():
                                style += "font-style:italic;"
                            
                            left_pct = (span["origin"][0] / width) * 100
                            top_pct = (span["origin"][1] / height) * 100
                            
                            span_html = (
                                f'<span style="position:absolute;'
                                f'left:{left_pct:.2f}%;top:{top_pct:.2f}%;'
                                f'{style}font-size:{span["size"]}pt;">'
                                f'{self._escape_html(span["text"])}</span>'
                            )
                            line_html.append(span_html)
                        block_html.append("".join(line_html))
                    page_html.append("".join(block_html))
            
            html_content.append(
                f'<div class="page" style="position:relative;width:100%;height:{height}px;">'
                f'{"".join(page_html)}</div>'
            )
        
        doc.close()
        
        full_html = (
            '<!DOCTYPE html><html><head><meta charset="UTF-8">'
            '<meta name="viewport" content="width=device-width, initial-scale=1.0">'
            '<style>.page { margin-bottom: 20px; }</style></head><body>'
            + "\n".join(html_content) + 
            '</body></html>'
        )
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_html)
        
        self._post_process_html(output_path, preserve_layout=True)
    
    def _convert_with_pdfplumber(self, pdf_path, output_path):
        """Conversion using pdfplumber for advanced table and text extraction"""
        html_content = []
        
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                tables = page.extract_tables({
                    "vertical_strategy": "text", 
                    "horizontal_strategy": "text"
                })
                
                text = page.extract_text(
                    x_tolerance=3,
                    y_tolerance=3,
                    layout=False
                )
                
                table_html = ""
                for table in tables:
                    df = pd.DataFrame(table)
                    table_html += df.to_html(index=False, header=False)
                
                page_content = f"<div class='page'><h2>Page {page_num}</h2>"
                if table_html:
                    page_content += f"<div class='tables'>{table_html}</div>"
                if text:
                    page_content += f"<div class='text'><pre>{self._escape_html(text)}</pre></div>"
                page_content += "</div>"
                
                html_content.append(page_content)
        
        full_html = (
            '<!DOCTYPE html><html><head><meta charset="UTF-8">'
            '<meta name="viewport" content="width=device-width, initial-scale=1.0">'
            '<style>'
            'body { font-family: Arial, sans-serif; line-height: 1.6; }'
            '.page { margin-bottom: 30px; border-bottom: 1px solid #eee; padding-bottom: 20px; }'
            'table { border-collapse: collapse; width: 100%; margin: 10px 0; }'
            'td, th { border: 1px solid #ddd; padding: 8px; }'
            'pre { white-space: pre-wrap; font-family: inherit; }'
            '</style></head><body>'
            + "\n".join(html_content) + 
            '</body></html>'
        )
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_html)
        
        self._post_process_html(output_path, preserve_layout=False)
    
    def _generate_clean_html(self, text_blocks):
        """Generate clean HTML from structured text blocks"""
        html_template = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Document</title>
    <style>
        :root {{
            font-size: 16px;
            color: #222;
        }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, 
                         Helvetica, Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 2rem;
        }}
        h1, h2, h3, h4 {{
            color: #111;
            line-height: 1.3;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }}
        h1 {{ font-size: 1.8rem; border-bottom: 1px solid #eee; padding-bottom: 0.3em; }}
        h2 {{ font-size: 1.5rem; }}
        h3 {{ font-size: 1.3rem; }}
        h4 {{ font-size: 1.1rem; }}
        p, li {{
            margin: 0 0 1rem 0;
        }}
        strong, b {{ font-weight: 600; }}
        em, i {{ font-style: italic; }}
        .page-break {{
            display: block;
            height: 0;
            border-top: 1px dashed #ccc;
            margin: 2rem 0;
        }}
        @media print {{
            body {{ padding: 0; }}
            .page-break {{ page-break-after: always; }}
        }}
    </style>
</head>
<body>
{content}
</body>
</html>"""
        
        current_page = 1
        html_content = []
        current_paragraph = []
        
        sorted_blocks = sorted(text_blocks, key=lambda x: (x['page'], x['y0'], x['x0']))
        
        for i, block in enumerate(sorted_blocks):
            if block['page'] > current_page:
                if current_paragraph:
                    html_content.append(f"<p>{' '.join(current_paragraph)}</p>")
                    current_paragraph = []
                html_content.append('<hr class="page-break">')
                current_page = block['page']
            
            text = self._escape_html(block['text'])
            
            if block['font_size'] and block['font_size'] > 14:
                if current_paragraph:
                    html_content.append(f"<p>{' '.join(current_paragraph)}</p>")
                    current_paragraph = []
                html_content.append(f'<h2>{text}</h2>')
            elif block['font_size'] and block['font_size'] > 12:
                if current_paragraph:
                    html_content.append(f"<p>{' '.join(current_paragraph)}</p>")
                    current_paragraph = []
                html_content.append(f'<h3>{text}</h3>')
            else:
                if block['bold'] and block['italic']:
                    text = f'<strong><em>{text}</em></strong>'
                elif block['bold']:
                    text = f'<strong>{text}</strong>'
                elif block['italic']:
                    text = f'<em>{text}</em>'
                
                current_paragraph.append(text)
        
        if current_paragraph:
            html_content.append(f"<p>{' '.join(current_paragraph)}</p>")
        
        return html_template.format(content='\n'.join(html_content))
    
    def _post_process_html(self, html_path, preserve_layout=False):
        """Enhance the HTML output quality with professional styling"""
        with open(html_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        
        meta = soup.new_tag('meta', attrs={
            'name': 'viewport',
            'content': 'width=device-width, initial-scale=1.0'
        })
        soup.head.insert(0, meta)
        
        if not soup.title:
            title = soup.new_tag('title')
            title.string = 'Converted PDF Document'
            soup.head.append(title)
        
        style = soup.new_tag('style')
        if preserve_layout:
            style.string = """
            body {
                font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, 
                            Helvetica, Arial, sans-serif;
                line-height: 1.5;
                color: #333;
                max-width: 100%;
                margin: 0;
                padding: 20px;
                background-color: #fff;
            }
            .page {
                position: relative;
                background: white;
                margin-bottom: 20px;
                box-shadow: 0 2px 5px rgba(0,0,0,0.1);
                overflow: hidden;
            }
            .textbox {
                position: absolute;
                white-space: pre-wrap;
            }
            @media print {
                body { padding: 0; }
                .page { 
                    box-shadow: none;
                    margin: 0;
                    page-break-after: always;
                }
            }
            """
        else:
            style.string = """
            body {
                font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, 
                            Helvetica, Arial, sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
                background-color: #fff;
            }
            h1, h2, h3, h4 {
                margin-top: 1.8em;
                margin-bottom: 0.6em;
                line-height: 1.3;
                color: #111;
            }
            h1 { 
                font-size: 2rem;
                border-bottom: 1px solid #eee;
                padding-bottom: 0.3em;
            }
            p, li {
                margin: 0 0 1.2em 0;
            }
            strong, b { 
                font-weight: 600;
                color: #000;
            }
            @media (max-width: 768px) {
                body {
                    padding: 15px;
                    font-size: 15px;
                }
            }
            @media print {
                body { 
                    padding: 0;
                    max-width: 100%;
                    font-size: 12pt;
                }
                h1, h2, h3 {
                    page-break-after: avoid;
                }
                p, li {
                    page-break-inside: avoid;
                }
            }
            """
        
        soup.head.append(style)
        
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(str(soup))
    
    def _save_uploaded_file(self):
        """Save uploaded file to temporary location with validation"""
        pdf_path = self.temp_dir / "input.pdf"
        
        if hasattr(self.file_obj, 'read'):
            # File-like object
            with open(pdf_path, 'wb+') as f:
                if hasattr(self.file_obj, 'chunks'):
                    for chunk in self.file_obj.chunks():
                        f.write(chunk)
                else:
                    f.write(self.file_obj.read())
        else:
            # Path or string
            with open(pdf_path, 'wb+') as f:
                with default_storage.open(self.file_obj, 'rb') as source_file:
                    f.write(source_file.read())
        
        # Verify PDF
        try:
            with fitz.open(pdf_path) as doc:
                if not doc.is_pdf:
                    raise PdfConversionError("Invalid PDF file")
        except Exception as e:
            raise PdfConversionError(f"PDF validation failed: {str(e)}")
        
        return pdf_path
    
    def _get_output_path(self, ext):
        """Generate output path in temp directory"""
        return self.temp_dir / f"output.{ext}"
    
    def _validate_output(self, output_path):
        """Validate the conversion output meets quality standards"""
        if not output_path.exists():
            return False
        
        if output_path.stat().st_size == 0:
            return False
        
        with open(output_path, 'r', encoding='utf-8') as f:
            content = f.read()
            if len(content) < 100 and '<body' in content:
                return False
            if '</body>' not in content:
                return False
        
        return True
    
    def _escape_html(self, text):
        """Escape HTML special characters"""
        return (text.replace('&', '&amp;')
                   .replace('<', '&lt;')
                   .replace('>', '&gt;')
                   .replace('"', '&quot;')
                   .replace("'", '&#39;'))
    
    def cleanup(self):
        """Clean up temporary files"""
        if self.temp_dir and self.temp_dir.exists():
            try:
                shutil.rmtree(self.temp_dir)
            except Exception as e:
                logger.error(f"Failed to clean up temp directory {self.temp_dir}: {str(e)}")
